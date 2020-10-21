import os
import re
import tkinter as tk
from io import StringIO
from tkinter import messagebox
from tkinter.filedialog import askopenfilename, asksaveasfilename, askopenfilenames

from PyPDF2 import PdfFileMerger
from docx import Document
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser

from constant import TEXT_FILES, TEXT_EXTENSION, PDF_FILES, PDF_FILES_EXTENSION, ALL_MS_WORD_EXTENSION, MS_WORD, \
    ALL_TEXT_EXTENSION


class FileUtil:

    def open_file(self, window, txt_edit):
        filepath = askopenfilename(
            filetypes=[(PDF_FILES, PDF_FILES_EXTENSION), (TEXT_FILES, TEXT_EXTENSION)]
        )
        if not filepath:
            return

        filename, file_extension = os.path.splitext(filepath)
        txt_edit.delete(1.0, tk.END)
        if file_extension == TEXT_EXTENSION:
            with open(filepath, "r") as input_file:
                txt_edit.insert(tk.END, input_file.read())
        elif file_extension == PDF_FILES_EXTENSION:
            self.convert_pdf_to_string(txt_edit, filepath)
        else:
            print("Not supported yet")

        window.title(f"Simple Text Editor - {filepath}")

    @classmethod
    def convert_pdf_to_string(self, txt_edit, path):
        output_string = StringIO()
        txt_edit.delete(1.0, tk.END)
        with open(path, 'rb') as in_file:
            parser = PDFParser(in_file)
            doc = PDFDocument(parser)
            rsrcmgr = PDFResourceManager()
            interpreter = PDFPageInterpreter(rsrcmgr, TextConverter(rsrcmgr, output_string, laparams=LAParams()))
            pages = PDFPage.create_pages(doc)
            for page in pages:
                interpreter.process_page(page)

            txt_edit.insert(tk.END, output_string.getvalue())

    def save_as(self, txt_edit):
        """Save the current file as a new pdf file."""
        filepath = asksaveasfilename(
            filetypes=[(MS_WORD, ALL_MS_WORD_EXTENSION), (TEXT_FILES, ALL_TEXT_EXTENSION)],
        )
        if not filepath:
            return

        document = Document()
        text = txt_edit.get(1.0, tk.END)
        document.add_paragraph(re.sub(r'\x0c', ' ', text))
        document.save(filepath)
        self.info('File saved successfully')

    def clear_all(self, txt_edit):
        txt_edit.delete(1.0, tk.END)

    def merge_pdf(self, txt_edit):
        files = askopenfilenames(
            filetypes=[(PDF_FILES, PDF_FILES_EXTENSION)]
        )
        if not files:
            return
        self.clear_all(txt_edit)
        pdf_merge = PdfFileMerger()
        for path in files:
            txt_edit.insert(tk.END, "Importing " + path + " ......\n")
            pdf_merge.append(path)

        filepath = asksaveasfilename(
            filetypes=[(PDF_FILES, PDF_FILES_EXTENSION)],
        )
        filename, file_extension = os.path.splitext(filepath)
        if file_extension != PDF_FILES_EXTENSION:
            filepath += PDF_FILES_EXTENSION
        if not files:
            return
        with open(filepath, "wb") as merge_file:
            pdf_merge.write(merge_file)
        msg = "Files are merged into " + filepath
        txt_edit.insert(tk.END, "\n\n " + msg)
        self.info(msg)

    def info(self, msg):
        messagebox.showinfo(title='Information....', message=msg)
